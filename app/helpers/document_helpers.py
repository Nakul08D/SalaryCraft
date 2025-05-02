import re
import os
import subprocess
from typing import Dict
from docx import Document
from docx.oxml.shared import qn


def replace_placeholders(doc: Document, data: Dict[str, str]) -> Document:
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for paragraph in all_paragraphs:
        replace_in_paragraph(paragraph, data)

    return doc

def replace_in_paragraph(paragraph, data):
    full_text = ''.join(run.text for run in paragraph.runs)
    placeholders = re.findall(r"{(.*?)}", full_text)
    if not placeholders:
        return

    original_runs = [(run.text, run) for run in paragraph.runs]
    paragraph.clear()

    for key in placeholders:
        placeholder = f"{{{key}}}"
        matched_key = next((k for k in data.keys() if k.lower() == key.lower()), None)
        value = data.get(matched_key, "") if matched_key else ""
        if matched_key and matched_key.lower() == "name" and value:
            value = value.upper()

        replacement = str(value) if value not in (None, "") else ""
        full_text = full_text.replace(placeholder, replacement)

    if original_runs:
        new_run = paragraph.add_run(full_text)
        copy_run_formatting(new_run, original_runs[0][1])
        for _, run in original_runs:
            copy_run_images(new_run, run)

def copy_run_formatting(new_run, original_run):
    new_run.bold = original_run.bold
    new_run.italic = original_run.italic
    new_run.underline = original_run.underline
    new_run.font.name = original_run.font.name
    new_run.font.size = original_run.font.size
    if original_run.font.color.rgb:
        new_run.font.color.rgb = original_run.font.color.rgb

def copy_run_images(new_run, original_run):
    for element in original_run.element:
        if element.tag.endswith('drawing'):
            new_run._r.append(element)

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        subprocess.run(
            [
                r"C:\Program Files\LibreOffice\program\soffice.exe",  
                "--headless", "--convert-to", "pdf",
                docx_path, "--outdir", os.path.dirname(pdf_path)
            ],
            check=True
        )
        return pdf_path
    except subprocess.CalledProcessError as e:
        raise Exception(f"Error converting DOCX to PDF: {str(e)}")